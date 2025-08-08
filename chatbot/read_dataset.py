#!/usr/bin/env python3
"""
Script to read and display the contents of Dataset_kmutnb.docx
"""

import os
import sys
from typing import Optional, List
from docx import Document


class DocumentReader:
    """Class for reading and processing Word documents"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
    
    def validate_file(self) -> bool:
        """Validate that the file exists and is accessible"""
        if not os.path.exists(self.file_path):
            print(f"Error: File {self.file_path} not found!")
            return False
        
        if not self.file_path.lower().endswith('.docx'):
            print(f"Error: File {self.file_path} is not a .docx file!")
            return False
        
        return True
    
    def read_paragraphs(self, doc: Document) -> List[str]:
        """Extract text from document paragraphs"""
        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        return content
    
    def read_tables(self, doc: Document) -> List[str]:
        """Extract text from document tables"""
        content = []
        for i, table in enumerate(doc.tables, 1):
            content.append(f"\n--- TABLE {i} ---")
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                content.append(" | ".join(row_data))
            content.append(f"--- END TABLE {i} ---\n")
        return content
    
    def read_document(self) -> Optional[str]:
        """Read and return the complete contents of the document"""
        if not self.validate_file():
            return None
        
        try:
            doc = Document(self.file_path)
            content = []
            
            # Read paragraphs
            content.extend(self.read_paragraphs(doc))
            
            # Read tables
            content.extend(self.read_tables(doc))
            
            return '\n'.join(content)
        
        except Exception as e:
            print(f"Error reading file: {e}")
            return None


def main():
    """Main function to read and display document content"""
    # Default file path
    file_path = "Dataset_kmutnb.docx"
    
    # Allow command line argument for file path
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
    
    print(f"Reading file: {file_path}")
    print("=" * 50)
    
    # Create document reader and read content
    reader = DocumentReader(file_path)
    content = reader.read_document()
    
    if content:
        print(content)
    else:
        print("Failed to read the file.")
        sys.exit(1)


if __name__ == "__main__":
    main() 