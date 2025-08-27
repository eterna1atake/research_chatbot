import os
import re
from typing import Optional, List, Dict, Tuple
from docx import Document
import fitz  # PyMuPDF
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from heapq import nlargest
from collections import Counter
import hashlib
import json

# OCR support
try:
    import pytesseract
    from PIL import Image
    import io
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# ‡∏î‡∏≤‡∏ß‡∏ô‡πå‡πÇ‡∏´‡∏•‡∏î NLTK data
try:
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)
    nltk.download('punkt_tab', quiet=True)
except:
    pass

class EnhancedDocumentReader:
    def __init__(self, file_path: str, use_ocr: bool = False, expert_role: str = ""):
        self.file_path = file_path
        self.content_cache = {}
        self.metadata = {}
        self.sections = {}
        self.keywords = set()
        self.use_ocr = use_ocr and OCR_AVAILABLE
        self.expert_role = expert_role
        
    def validate_file(self) -> bool:
        """‡∏ï‡∏£‡∏ß‡∏à‡∏™‡∏≠‡∏ö‡πÑ‡∏ü‡∏•‡πå‡∏≠‡∏¢‡πà‡∏≤‡∏á‡∏•‡∏∞‡πÄ‡∏≠‡∏µ‡∏¢‡∏î"""
        if not os.path.exists(self.file_path):
            print(f"Error: File {self.file_path} not found!")
            return False

        # ‡πÄ‡∏ä‡πá‡∏Ñ‡∏Ç‡∏ô‡∏≤‡∏î‡πÑ‡∏ü‡∏•‡πå
        try:
            file_size = os.path.getsize(self.file_path)
            if file_size == 0:
                print(f"Error: File {self.file_path} is empty!")
                return False
            elif file_size > 100 * 1024 * 1024:  # 100MB
                print(f"Warning: File {self.file_path} is very large ({file_size/(1024*1024):.1f}MB)")
        except:
            pass

        supported_extensions = ['.docx', '.pdf', '.txt', '.doc']
        if not any(self.file_path.lower().endswith(ext) for ext in supported_extensions):
            print(f"Error: Unsupported file type {self.file_path}!")
            return False

        return True

    def extract_keywords(self, text: str, num_keywords: int = 50) -> List[str]:
        """‡∏™‡∏Å‡∏±‡∏î‡∏Ñ‡∏≥‡∏™‡∏≥‡∏Ñ‡∏±‡∏ç‡∏à‡∏≤‡∏Å‡∏Ç‡πâ‡∏≠‡∏Ñ‡∏ß‡∏≤‡∏°"""
        try:
            # ‡∏ó‡∏≥‡∏Ñ‡∏ß‡∏≤‡∏°‡∏™‡∏∞‡∏≠‡∏≤‡∏î‡∏Ç‡πâ‡∏≠‡∏Ñ‡∏ß‡∏≤‡∏°
            text = re.sub(r'[^\w\s]', ' ', text.lower())
            words = word_tokenize(text)
            
            # ‡∏Å‡∏£‡∏≠‡∏á‡∏Ñ‡∏≥‡∏ó‡∏µ‡πà‡πÑ‡∏°‡πà‡∏ï‡πâ‡∏≠‡∏á‡∏Å‡∏≤‡∏£
            try:
                stop_words = set(stopwords.words('english'))
                stop_words.update(['‡πÅ‡∏•‡∏∞', '‡∏´‡∏£‡∏∑‡∏≠', '‡∏ó‡∏µ‡πà', '‡πÄ‡∏õ‡πá‡∏ô', '‡πÉ‡∏ô', '‡∏à‡∏≤‡∏Å', '‡∏Ç‡∏≠‡∏á', '‡∏Å‡∏≤‡∏£', '‡πÑ‡∏î‡πâ', '‡∏°‡∏µ', '‡πÉ‡∏´‡πâ'])
            except:
                stop_words = set()
                
            # ‡∏Å‡∏£‡∏≠‡∏á‡∏Ñ‡∏≥‡∏ó‡∏µ‡πà‡∏°‡∏µ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏¢‡∏≤‡∏ß‡πÄ‡∏´‡∏°‡∏≤‡∏∞‡∏™‡∏°
            filtered_words = [
                word for word in words 
                if len(word) > 2 and word not in stop_words and not word.isdigit()
            ]
            
            # ‡∏ô‡∏±‡∏ö‡∏Ñ‡∏ß‡∏≤‡∏°‡∏ñ‡∏µ‡πà‡πÅ‡∏•‡∏∞‡πÄ‡∏•‡∏∑‡∏≠‡∏Å‡∏Ñ‡∏≥‡∏™‡∏≥‡∏Ñ‡∏±‡∏ç
            word_freq = Counter(filtered_words)
            keywords = [word for word, freq in word_freq.most_common(num_keywords)]
            
            self.keywords.update(keywords)
            return keywords
            
        except Exception as e:
            print(f"Error extracting keywords: {e}")
            return []

    def segment_text(self, text: str) -> Dict[str, str]:
        """‡πÅ‡∏ö‡πà‡∏á‡∏Ç‡πâ‡∏≠‡∏Ñ‡∏ß‡∏≤‡∏°‡πÄ‡∏õ‡πá‡∏ô‡∏™‡πà‡∏ß‡∏ô‡πÜ ‡∏ï‡∏≤‡∏°‡∏´‡∏±‡∏ß‡∏Ç‡πâ‡∏≠"""
        sections = {}
        current_section = "introduction"
        current_content = []
        
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # ‡∏ï‡∏£‡∏ß‡∏à‡∏´‡∏≤‡∏´‡∏±‡∏ß‡∏Ç‡πâ‡∏≠ (‡∏ö‡∏£‡∏£‡∏ó‡∏±‡∏î‡∏ó‡∏µ‡πà‡πÄ‡∏õ‡πá‡∏ô‡∏ï‡∏±‡∏ß‡∏û‡∏¥‡∏°‡∏û‡πå‡πÉ‡∏´‡∏ç‡πà‡∏´‡∏£‡∏∑‡∏≠‡∏°‡∏µ‡∏£‡∏π‡∏õ‡πÅ‡∏ö‡∏ö‡∏û‡∏¥‡πÄ‡∏®‡∏©)
            if (len(line) < 100 and 
                (line.isupper() or 
                 re.match(r'^[\d\.\s]*[A-Za-z‡∏Å-‡πô]+', line) or
                 line.startswith('‡∏ö‡∏ó') or line.startswith('‡∏´‡∏°‡∏ß‡∏î') or
                 line.startswith('‡∏™‡πà‡∏ß‡∏ô‡∏ó‡∏µ‡πà') or line.startswith('Chapter'))):
                
                # ‡∏ö‡∏±‡∏ô‡∏ó‡∏∂‡∏Å‡∏™‡πà‡∏ß‡∏ô‡∏ó‡∏µ‡πà‡πÅ‡∏•‡πâ‡∏ß
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                
                # ‡πÄ‡∏£‡∏¥‡πà‡∏°‡∏™‡πà‡∏ß‡∏ô‡πÉ‡∏´‡∏°‡πà
                current_section = line.lower().replace(' ', '_')
                current_content = []
            else:
                current_content.append(line)
        
        # ‡∏ö‡∏±‡∏ô‡∏ó‡∏∂‡∏Å‡∏™‡πà‡∏ß‡∏ô‡∏™‡∏∏‡∏î‡∏ó‡πâ‡∏≤‡∏¢
        if current_content:
            sections[current_section] = '\n'.join(current_content)
            
        self.sections = sections
        return sections

    def enhanced_clean_text(self, text: str) -> str:
        """‡∏ó‡∏≥‡∏Ñ‡∏ß‡∏≤‡∏°‡∏™‡∏∞‡∏≠‡∏≤‡∏î‡∏Ç‡πâ‡∏≠‡∏Ñ‡∏ß‡∏≤‡∏°‡∏≠‡∏¢‡πà‡∏≤‡∏á‡∏•‡∏∞‡πÄ‡∏≠‡∏µ‡∏¢‡∏î"""
        if not text:
            return ""
        
        # ‡∏•‡∏ö‡∏≠‡∏±‡∏Å‡∏Ç‡∏£‡∏∞‡∏û‡∏¥‡πÄ‡∏®‡∏©‡∏ó‡∏µ‡πà‡πÑ‡∏°‡πà‡∏ï‡πâ‡∏≠‡∏á‡∏Å‡∏≤‡∏£
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', '', text)
        
        # ‡∏õ‡∏£‡∏±‡∏ö‡∏õ‡∏£‡∏∏‡∏á‡∏Å‡∏≤‡∏£‡∏Ç‡∏∂‡πâ‡∏ô‡∏ö‡∏£‡∏£‡∏ó‡∏±‡∏î‡πÉ‡∏´‡∏°‡πà
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # ‡∏•‡∏î‡∏ö‡∏£‡∏£‡∏ó‡∏±‡∏î‡∏ß‡πà‡∏≤‡∏á‡∏ã‡πâ‡∏≥
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            
            # ‡∏Å‡∏£‡∏≠‡∏á‡∏ö‡∏£‡∏£‡∏ó‡∏±‡∏î‡∏ó‡∏µ‡πà‡πÑ‡∏°‡πà‡∏°‡∏µ‡∏õ‡∏£‡∏∞‡πÇ‡∏¢‡∏ä‡∏ô‡πå
            if (len(line) < 3 or 
                line.isdigit() or 
                re.match(r'^[\s\-_=]+$', line) or
                line.count('.') > len(line) * 0.3):  # ‡∏ö‡∏£‡∏£‡∏ó‡∏±‡∏î‡∏ó‡∏µ‡πà‡∏°‡∏µ‡∏à‡∏∏‡∏î‡πÄ‡∏¢‡∏≠‡∏∞‡πÄ‡∏Å‡∏¥‡∏ô‡πÑ‡∏õ
                continue
                
            # ‡∏õ‡∏£‡∏±‡∏ö‡∏õ‡∏£‡∏∏‡∏á‡∏ä‡πà‡∏≠‡∏á‡∏ß‡πà‡∏≤‡∏á
            line = re.sub(r'\s+', ' ', line)
            
            # ‡∏£‡∏ß‡∏°‡∏ö‡∏£‡∏£‡∏ó‡∏±‡∏î‡∏ó‡∏µ‡πà‡∏Ç‡∏≤‡∏î‡πÜ ‡∏´‡∏≤‡∏¢‡πÜ
            if (len(cleaned_lines) > 0 and 
                not line.endswith(('.', '!', '?', ':', ';')) and
                len(line) < 80 and
                not line[0].isupper()):
                cleaned_lines[-1] += ' ' + line
            else:
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)

    def read_txt_file(self) -> str:
        """‡∏≠‡πà‡∏≤‡∏ô‡πÑ‡∏ü‡∏•‡πå txt with multiple encodings"""
        encodings = ['utf-8', 'utf-8-sig', 'cp874', 'tis-620', 'iso-8859-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(self.file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                    return self.enhanced_clean_text(content)
            except UnicodeDecodeError:
                continue
            except Exception as e:
                print(f"Error with encoding {encoding}: {e}")
                continue
                
        return "Error: Cannot read text file with any supported encoding"

    def read_docx_advanced(self) -> str:
        """‡∏≠‡πà‡∏≤‡∏ô DOCX ‡∏≠‡∏¢‡πà‡∏≤‡∏á‡∏•‡∏∞‡πÄ‡∏≠‡∏µ‡∏¢‡∏î"""
        try:
            doc = Document(self.file_path)
            content = []
            
            # ‡∏≠‡πà‡∏≤‡∏ô paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    # ‡πÄ‡∏ä‡πá‡∏Ñ style ‡∏Ç‡∏≠‡∏á paragraph
                    style_name = para.style.name if para.style else ""
                    text = para.text.strip()
                    
                    # ‡πÄ‡∏û‡∏¥‡πà‡∏° marker ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡∏´‡∏±‡∏ß‡∏Ç‡πâ‡∏≠
                    if 'heading' in style_name.lower() or para.style.font.bold:
                        text = f"\n=== {text} ===\n"
                    
                    content.append(text)
            
            # ‡∏≠‡πà‡∏≤‡∏ô tables ‡∏≠‡∏¢‡πà‡∏≤‡∏á‡∏•‡∏∞‡πÄ‡∏≠‡∏µ‡∏¢‡∏î
            for i, table in enumerate(doc.tables, 1):
                content.append(f"\n--- ‡∏ï‡∏≤‡∏£‡∏≤‡∏á {i} ---")
                
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_data.append(cell_text)
                    
                    if row_data:  # ‡∏°‡∏µ‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡πÉ‡∏ô‡πÅ‡∏ñ‡∏ß
                        if row_idx == 0:  # ‡πÅ‡∏ñ‡∏ß‡πÅ‡∏£‡∏Å‡∏≠‡∏≤‡∏à‡πÄ‡∏õ‡πá‡∏ô‡∏´‡∏±‡∏ß‡∏ï‡∏≤‡∏£‡∏≤‡∏á
                            content.append("‡∏´‡∏±‡∏ß‡∏ï‡∏≤‡∏£‡∏≤‡∏á: " + " | ".join(row_data))
                        else:
                            content.append(" | ".join(row_data))
                
                content.append(f"--- ‡∏à‡∏ö‡∏ï‡∏≤‡∏£‡∏≤‡∏á {i} ---\n")
            
            # ‡∏≠‡πà‡∏≤‡∏ô headers ‡πÅ‡∏•‡∏∞ footers
            for section in doc.sections:
                if section.header:
                    header_text = ""
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            header_text += para.text.strip() + " "
                    if header_text:
                        content.insert(0, f"=== HEADER: {header_text.strip()} ===\n")
            
            full_content = '\n'.join(content)
            return self.enhanced_clean_text(full_content)
            
        except Exception as e:
            return f"Error reading DOCX: {str(e)}"

    def extract_text_with_ocr(self, page) -> str:
        """‡πÉ‡∏ä‡πâ OCR ‡πÄ‡∏û‡∏∑‡πà‡∏≠‡∏î‡∏∂‡∏á‡∏Ç‡πâ‡∏≠‡∏Ñ‡∏ß‡∏≤‡∏°‡∏à‡∏≤‡∏Å‡∏´‡∏ô‡πâ‡∏≤ PDF"""
        try:
            # ‡πÅ‡∏õ‡∏•‡∏á‡∏´‡∏ô‡πâ‡∏≤‡πÄ‡∏õ‡πá‡∏ô‡∏†‡∏≤‡∏û
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # ‡πÄ‡∏û‡∏¥‡πà‡∏°‡∏Ñ‡∏ß‡∏≤‡∏°‡∏•‡∏∞‡πÄ‡∏≠‡∏µ‡∏¢‡∏î
            img_data = pix.tobytes("png")
            
            # ‡πÄ‡∏õ‡∏¥‡∏î‡∏†‡∏≤‡∏û‡∏î‡πâ‡∏ß‡∏¢ PIL
            img = Image.open(io.BytesIO(img_data))
            
            # ‡πÉ‡∏ä‡πâ OCR
            text = pytesseract.image_to_string(img, lang='tha+eng')
            
            return text.strip()
        except Exception as e:
            print(f"OCR Error: {e}")
            return ""

    def read_pdf_advanced(self) -> str:
        """‡∏≠‡πà‡∏≤‡∏ô PDF ‡∏≠‡∏¢‡πà‡∏≤‡∏á‡∏•‡∏∞‡πÄ‡∏≠‡∏µ‡∏¢‡∏î‡πÅ‡∏•‡∏∞‡πÅ‡∏°‡πà‡∏ô‡∏¢‡∏≥ ‡∏û‡∏£‡πâ‡∏≠‡∏° OCR support"""
        try:
            doc = fitz.open(self.file_path)
            full_content = []
            
            # ‡πÄ‡∏Å‡πá‡∏ö‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏• metadata
            metadata = doc.metadata
            if metadata:
                self.metadata = {
                    'title': metadata.get('title', ''),
                    'author': metadata.get('author', ''),
                    'subject': metadata.get('subject', ''),
                    'pages': len(doc)
                }
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # ‡∏•‡∏≠‡∏á‡∏î‡∏∂‡∏á‡∏Ç‡πâ‡∏≠‡∏Ñ‡∏ß‡∏≤‡∏°‡πÅ‡∏ö‡∏ö‡∏õ‡∏Å‡∏ï‡∏¥‡∏Å‡πà‡∏≠‡∏ô
                text = page.get_text("text")
                
                # ‡∏ñ‡πâ‡∏≤‡πÑ‡∏°‡πà‡∏°‡∏µ‡∏Ç‡πâ‡∏≠‡∏Ñ‡∏ß‡∏≤‡∏°‡∏´‡∏£‡∏∑‡∏≠‡∏Ç‡πâ‡∏≠‡∏Ñ‡∏ß‡∏≤‡∏°‡∏ô‡πâ‡∏≠‡∏¢‡πÄ‡∏Å‡∏¥‡∏ô‡πÑ‡∏õ ‡πÅ‡∏•‡∏∞‡πÄ‡∏õ‡∏¥‡∏î‡πÉ‡∏ä‡πâ OCR
                if (not text.strip() or len(text.strip()) < 50) and self.use_ocr:
                    print(f"Using OCR for page {page_num + 1}")
                    text = self.extract_text_with_ocr(page)
                    if text:
                        text = f"[OCR] {text}"
                
                # ‡∏ñ‡πâ‡∏≤‡∏¢‡∏±‡∏á‡πÑ‡∏°‡πà‡∏°‡∏µ‡∏Ç‡πâ‡∏≠‡∏Ñ‡∏ß‡∏≤‡∏° ‡∏•‡∏≠‡∏á‡∏ß‡∏¥‡∏ò‡∏µ‡∏≠‡∏∑‡πà‡∏ô
                if not text.strip():
                    try:
                        blocks = page.get_text("dict")["blocks"]
                        text_blocks = []
                        
                        for block in blocks:
                            if "lines" in block:  # text block
                                for line in block["lines"]:
                                    line_text = ""
                                    for span in line["spans"]:
                                        line_text += span["text"]
                                    if line_text.strip():
                                        text_blocks.append(line_text.strip())
                        
                        text = '\n'.join(text_blocks)
                    except:
                        text = f"[‡∏´‡∏ô‡πâ‡∏≤ {page_num + 1}: ‡πÑ‡∏°‡πà‡∏™‡∏≤‡∏°‡∏≤‡∏£‡∏ñ‡∏î‡∏∂‡∏á‡∏Ç‡πâ‡∏≠‡∏Ñ‡∏ß‡∏≤‡∏°‡πÑ‡∏î‡πâ]"
                
                if text and text.strip():
                    cleaned_text = self.enhanced_clean_text(text)
                    if cleaned_text:
                        full_content.append(f"\n=== ‡∏´‡∏ô‡πâ‡∏≤ {page_num + 1} ===")
                        full_content.append(cleaned_text)
                        full_content.append("=" * 50)
            
            doc.close()
            
            final_content = '\n'.join(full_content)
            return final_content.strip() if final_content else "‡πÑ‡∏°‡πà‡∏™‡∏≤‡∏°‡∏≤‡∏£‡∏ñ‡∏î‡∏∂‡∏á‡∏Ç‡πâ‡∏≠‡∏Ñ‡∏ß‡∏≤‡∏°‡∏à‡∏≤‡∏Å PDF ‡πÑ‡∏î‡πâ"
            
        except Exception as e:
            return f"Error reading PDF: {str(e)}"

    def create_content_index(self, content: str) -> Dict[str, List[Tuple[str, int]]]:
        """‡∏™‡∏£‡πâ‡∏≤‡∏á index ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡∏Ñ‡πâ‡∏ô‡∏´‡∏≤"""
        index = {}
        lines = content.split('\n')
        
        for line_num, line in enumerate(lines):
            words = re.findall(r'\b\w+\b', line.lower())
            for word in words:
                if len(word) > 2:  # ‡πÄ‡∏â‡∏û‡∏≤‡∏∞‡∏Ñ‡∏≥‡∏ó‡∏µ‡πà‡∏°‡∏µ‡∏Ñ‡∏ß‡∏≤‡∏°‡∏¢‡∏≤‡∏ß‡∏°‡∏≤‡∏Å‡∏Å‡∏ß‡πà‡∏≤ 2
                    if word not in index:
                        index[word] = []
                    index[word].append((line.strip(), line_num))
        
        return index

    def smart_search(self, content: str, search_term: str, context_lines: int = 2) -> str:
        """‡∏Ñ‡πâ‡∏ô‡∏´‡∏≤‡πÅ‡∏ö‡∏ö‡∏≠‡∏±‡∏à‡∏â‡∏£‡∏¥‡∏¢‡∏∞"""
        if not content or not search_term:
            return "‡πÑ‡∏°‡πà‡∏°‡∏µ‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡∏Ñ‡πâ‡∏ô‡∏´‡∏≤"
        
        search_terms = search_term.lower().split()
        lines = content.split('\n')
        results = []
        
        for i, line in enumerate(lines):
            line_lower = line.lower()
            
            # ‡∏Ñ‡πâ‡∏ô‡∏´‡∏≤‡πÅ‡∏ö‡∏ö exact match
            if search_term.lower() in line_lower:
                context_start = max(0, i - context_lines)
                context_end = min(len(lines), i + context_lines + 1)
                context = lines[context_start:context_end]
                
                # ‡πÑ‡∏Æ‡πÑ‡∏•‡∏ï‡πå‡∏Ñ‡∏≥‡∏ó‡∏µ‡πà‡∏Ñ‡πâ‡∏ô‡∏´‡∏≤
                highlighted_line = line
                for term in search_terms:
                    highlighted_line = re.sub(
                        f'({re.escape(term)})', 
                        r'**\1**', 
                        highlighted_line, 
                        flags=re.IGNORECASE
                    )
                
                result = {
                    'line_number': i + 1,
                    'matched_line': highlighted_line,
                    'context': context,
                    'relevance_score': len([t for t in search_terms if t in line_lower])
                }
                results.append(result)
            
            # ‡∏Ñ‡πâ‡∏ô‡∏´‡∏≤‡πÅ‡∏ö‡∏ö partial match
            elif any(term in line_lower for term in search_terms):
                relevance = sum(1 for term in search_terms if term in line_lower)
                if relevance >= len(search_terms) * 0.5:  # ‡∏≠‡∏¢‡πà‡∏≤‡∏á‡∏ô‡πâ‡∏≠‡∏¢‡∏Ñ‡∏£‡∏∂‡πà‡∏á‡∏´‡∏ô‡∏∂‡πà‡∏á‡∏Ç‡∏≠‡∏á‡∏Ñ‡∏≥‡∏Ñ‡πâ‡∏ô‡∏´‡∏≤
                    context_start = max(0, i - context_lines)
                    context_end = min(len(lines), i + context_lines + 1)
                    context = lines[context_start:context_end]
                    
                    result = {
                        'line_number': i + 1,
                        'matched_line': line,
                        'context': context,
                        'relevance_score': relevance
                    }
                    results.append(result)
        
        # ‡πÄ‡∏£‡∏µ‡∏¢‡∏á‡∏ï‡∏≤‡∏° relevance score
        results.sort(key=lambda x: x['relevance_score'], reverse=True)
        
        if not results:
            # ‡∏•‡∏≠‡∏á‡∏Ñ‡πâ‡∏ô‡∏´‡∏≤‡∏Ñ‡∏≥‡∏ó‡∏µ‡πà‡∏Ñ‡∏•‡πâ‡∏≤‡∏¢‡∏Å‡∏±‡∏ô
            similar_results = []
            for i, line in enumerate(lines):
                similarity = 0
                for term in search_terms:
                    if any(term in word for word in line.lower().split()):
                        similarity += 1
                
                if similarity > 0:
                    similar_results.append((i + 1, line, similarity))
            
            if similar_results:
                similar_results.sort(key=lambda x: x[2], reverse=True)
                return f"‚ùì ‡πÑ‡∏°‡πà‡∏û‡∏ö‡∏Ñ‡∏≥‡∏ß‡πà‡∏≤ '{search_term}' ‡πÅ‡∏ï‡πà‡∏û‡∏ö‡∏Ñ‡∏≥‡∏ó‡∏µ‡πà‡∏Ñ‡∏•‡πâ‡∏≤‡∏¢‡∏Å‡∏±‡∏ô‡πÉ‡∏ô‡∏ö‡∏£‡∏£‡∏ó‡∏±‡∏î‡∏ó‡∏µ‡πà {similar_results[0][0]}: {similar_results[0][1][:100]}..."
            else:
                return f"‚ùå ‡πÑ‡∏°‡πà‡∏û‡∏ö‡∏Ñ‡∏≥‡∏ß‡πà‡∏≤ '{search_term}' ‡πÉ‡∏ô‡πÄ‡∏≠‡∏Å‡∏™‡∏≤‡∏£"
        
        # ‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏ú‡∏•‡∏•‡∏±‡∏û‡∏ò‡πå
        output = [f"üîç ‡∏û‡∏ö‡∏Ñ‡∏≥‡∏ß‡πà‡∏≤ '{search_term}' ‡πÉ‡∏ô‡πÄ‡∏≠‡∏Å‡∏™‡∏≤‡∏£ {len(results)} ‡∏ï‡∏≥‡πÅ‡∏´‡∏ô‡πà‡∏á:\n"]
        
        for i, result in enumerate(results[:10]):  # ‡πÅ‡∏™‡∏î‡∏á‡∏™‡∏π‡∏á‡∏™‡∏∏‡∏î 10 ‡∏ú‡∏•‡∏•‡∏±‡∏û‡∏ò‡πå
            output.append(f"\n=== ‡∏ú‡∏•‡∏•‡∏±‡∏û‡∏ò‡πå‡∏ó‡∏µ‡πà {i + 1} (‡∏ö‡∏£‡∏£‡∏ó‡∏±‡∏î {result['line_number']}) ===")
            output.append(f"üìç {result['matched_line']}")
            
            if result['context']:
                output.append("\nüìñ ‡∏ö‡∏£‡∏¥‡∏ö‡∏ó:")
                for ctx_line in result['context']:
                    if ctx_line.strip():
                        prefix = "‚û§ " if ctx_line == result['matched_line'] else "  "
                        output.append(f"{prefix}{ctx_line}")
        
        if len(results) > 10:
            output.append(f"\n... ‡πÅ‡∏•‡∏∞‡∏≠‡∏µ‡∏Å {len(results) - 10} ‡∏ú‡∏•‡∏•‡∏±‡∏û‡∏ò‡πå")
        
        return '\n'.join(output)

    def get_comprehensive_summary(self, max_chars: int = 20000) -> str:
        """‡∏™‡∏£‡πâ‡∏≤‡∏á‡∏™‡∏£‡∏∏‡∏õ‡πÄ‡∏≠‡∏Å‡∏™‡∏≤‡∏£‡∏≠‡∏¢‡πà‡∏≤‡∏á‡∏•‡∏∞‡πÄ‡∏≠‡∏µ‡∏¢‡∏î"""
        if not self.validate_file():
            return "Error: Could not read the dataset file."

        try:
            # ‡∏≠‡πà‡∏≤‡∏ô‡πÄ‡∏≠‡∏Å‡∏™‡∏≤‡∏£‡∏ï‡∏≤‡∏°‡∏õ‡∏£‡∏∞‡πÄ‡∏†‡∏ó
            if self.file_path.lower().endswith('.txt'):
                content = self.read_txt_file()
            elif self.file_path.lower().endswith('.docx'):
                content = self.read_docx_advanced()
            elif self.file_path.lower().endswith('.pdf'):
                content = self.read_pdf_advanced()
            else:
                return "Error: Unsupported file type"
            
            if not content or content.startswith("Error"):
                return content
            
            # ‡∏™‡∏£‡πâ‡∏≤‡∏á sections ‡πÅ‡∏•‡∏∞ keywords
            sections = self.segment_text(content)
            keywords = self.extract_keywords(content)
            
            # ‡∏™‡∏£‡πâ‡∏≤‡∏á index ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡∏Ñ‡πâ‡∏ô‡∏´‡∏≤
            content_index = self.create_content_index(content)
            
            # ‡∏ï‡∏±‡∏î‡πÄ‡∏ô‡∏∑‡πâ‡∏≠‡∏´‡∏≤‡∏ñ‡πâ‡∏≤‡∏¢‡∏≤‡∏ß‡πÄ‡∏Å‡∏¥‡∏ô‡πÑ‡∏õ
            original_length = len(content)
            if len(content) > max_chars:
                # ‡∏û‡∏¢‡∏≤‡∏¢‡∏≤‡∏°‡∏ï‡∏±‡∏î‡∏ó‡∏µ‡πà‡∏à‡∏∏‡∏î‡∏ó‡∏µ‡πà‡πÄ‡∏´‡∏°‡∏≤‡∏∞‡∏™‡∏°
                content = content[:max_chars]
                
                # ‡∏´‡∏≤‡∏à‡∏∏‡∏î‡∏ï‡∏±‡∏î‡∏ó‡∏µ‡πà‡∏î‡∏µ
                for cutoff in ['.', '\n\n', '\n', ' ']:
                    last_cutoff = content.rfind(cutoff)
                    if last_cutoff > max_chars * 0.8:
                        content = content[:last_cutoff + (1 if cutoff == '.' else 0)]
                        break
                
                content += f"\n\n[üìÑ ‡∏´‡∏°‡∏≤‡∏¢‡πÄ‡∏´‡∏ï‡∏∏: ‡πÄ‡∏ô‡∏∑‡πâ‡∏≠‡∏´‡∏≤‡∏ñ‡∏π‡∏Å‡∏ï‡∏±‡∏î‡πÄ‡∏û‡∏∑‡πà‡∏≠‡∏õ‡∏£‡∏∞‡∏´‡∏¢‡∏±‡∏î token ‡∏à‡∏≤‡∏Å {original_length:,} ‡πÄ‡∏õ‡πá‡∏ô {len(content):,} ‡∏ï‡∏±‡∏ß‡∏≠‡∏±‡∏Å‡∏©‡∏£]"
            
            # ‡πÄ‡∏û‡∏¥‡πà‡∏°‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏• metadata
            summary_parts = []
            
            # ‡πÄ‡∏û‡∏¥‡πà‡∏° expert role ‡∏ñ‡πâ‡∏≤‡∏°‡∏µ
            if self.expert_role:
                summary_parts.append(f"=== ‡∏ú‡∏π‡πâ‡πÄ‡∏ä‡∏µ‡πà‡∏¢‡∏ß‡∏ä‡∏≤‡∏ç: {self.expert_role} ===")
                summary_parts.append("")
            
            if self.metadata:
                summary_parts.append("=== ‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏•‡πÄ‡∏≠‡∏Å‡∏™‡∏≤‡∏£ ===")
                for key, value in self.metadata.items():
                    if value:
                        summary_parts.append(f"{key}: {value}")
                summary_parts.append("")
            
            if keywords:
                summary_parts.append("=== ‡∏Ñ‡∏≥‡∏™‡∏≥‡∏Ñ‡∏±‡∏ç (Top 20) ===")
                summary_parts.append(", ".join(keywords[:20]))
                summary_parts.append("")
            
            if len(sections) > 1:
                summary_parts.append("=== ‡∏´‡∏±‡∏ß‡∏Ç‡πâ‡∏≠‡πÉ‡∏ô‡πÄ‡∏≠‡∏Å‡∏™‡∏≤‡∏£ ===")
                for section_name in sections.keys():
                    summary_parts.append(f"- {section_name.replace('_', ' ').title()}")
                summary_parts.append("")
            
            summary_parts.append("=== ‡πÄ‡∏ô‡∏∑‡πâ‡∏≠‡∏´‡∏≤‡πÄ‡∏≠‡∏Å‡∏™‡∏≤‡∏£ ===")
            summary_parts.append(content)
            
            # ‡πÄ‡∏û‡∏¥‡πà‡∏°‡∏Ç‡πâ‡∏≠‡∏°‡∏π‡∏• OCR ‡∏ñ‡πâ‡∏≤‡πÉ‡∏ä‡πâ
            if self.use_ocr:
                summary_parts.append(f"\n\n[üîç ‡πÉ‡∏ä‡πâ OCR ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡∏Å‡∏≤‡∏£‡∏õ‡∏£‡∏∞‡∏°‡∏ß‡∏•‡∏ú‡∏• PDF]")
            
            summary_parts.append(f"\n\n[‚úÖ ‡πÄ‡∏≠‡∏Å‡∏™‡∏≤‡∏£ KMUTNB ‡∏ñ‡∏π‡∏Å‡πÇ‡∏´‡∏•‡∏î‡πÅ‡∏•‡∏∞‡∏õ‡∏£‡∏∞‡∏°‡∏ß‡∏•‡∏ú‡∏•‡πÄ‡∏£‡∏µ‡∏¢‡∏ö‡∏£‡πâ‡∏≠‡∏¢‡πÅ‡∏•‡πâ‡∏ß - ‡∏£‡∏ß‡∏° {len(content):,} ‡∏ï‡∏±‡∏ß‡∏≠‡∏±‡∏Å‡∏©‡∏£]")
            
            final_content = '\n'.join(summary_parts)
            
            # ‡πÄ‡∏Å‡πá‡∏ö‡πÑ‡∏ß‡πâ‡πÉ‡∏ô cache
            self.content_cache['full_content'] = final_content
            self.content_cache['sections'] = sections
            self.content_cache['keywords'] = keywords
            self.content_cache['index'] = content_index
            
            return final_content
                
        except Exception as e:
            return f"Error reading file: {str(e)}"

    def read_document(self) -> Optional[str]:
        """‡∏≠‡πà‡∏≤‡∏ô‡πÄ‡∏≠‡∏Å‡∏™‡∏≤‡∏£‡πÅ‡∏ö‡∏ö‡πÄ‡∏ï‡πá‡∏° (backward compatibility)"""
        return self.get_comprehensive_summary(max_chars=50000)

# ‡∏ü‡∏±‡∏á‡∏Å‡πå‡∏ä‡∏±‡∏ô‡∏´‡∏•‡∏±‡∏Å‡∏ó‡∏µ‡πà app.py ‡∏à‡∏∞‡πÄ‡∏£‡∏µ‡∏¢‡∏Å‡πÉ‡∏ä‡πâ
def get_kmutnb_summary(file_path: str, use_ocr: bool = False, expert_role: str = "") -> str:
    """‡∏ü‡∏±‡∏á‡∏Å‡πå‡∏ä‡∏±‡∏ô‡∏´‡∏•‡∏±‡∏Å‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö‡∏≠‡πà‡∏≤‡∏ô‡πÅ‡∏•‡∏∞‡∏™‡∏£‡∏∏‡∏õ‡πÄ‡∏≠‡∏Å‡∏™‡∏≤‡∏£ KMUTNB (‡πÅ‡∏ö‡∏ö‡πÉ‡∏´‡∏°‡πà)"""
    reader = EnhancedDocumentReader(file_path, use_ocr=use_ocr, expert_role=expert_role)
    return reader.get_comprehensive_summary()

def read_kmutnb_dataset(file_path: str, use_ocr: bool = False, expert_role: str = "") -> str:
    """‡∏≠‡πà‡∏≤‡∏ô‡πÄ‡∏≠‡∏Å‡∏™‡∏≤‡∏£ KMUTNB ‡πÅ‡∏ö‡∏ö‡πÄ‡∏ï‡πá‡∏° (‡πÅ‡∏ö‡∏ö‡πÉ‡∏´‡∏°‡πà)"""
    reader = EnhancedDocumentReader(file_path, use_ocr=use_ocr, expert_role=expert_role)
    content = reader.get_comprehensive_summary(max_chars=100000)
    if content is None or content.startswith("Error"):
        return "Error: Could not read the dataset file."
    return content

def search_in_document(file_path: str, search_term: str, use_ocr: bool = False) -> str:
    """‡∏Ñ‡πâ‡∏ô‡∏´‡∏≤‡∏Ñ‡∏≥‡πÉ‡∏ô‡πÄ‡∏≠‡∏Å‡∏™‡∏≤‡∏£ (‡πÅ‡∏ö‡∏ö‡πÉ‡∏´‡∏°‡πà)"""
    reader = EnhancedDocumentReader(file_path, use_ocr=use_ocr)
    content = reader.read_document()
    
    if not content or content.startswith("Error"):
        return "Error: Could not read the document."
    
    return reader.smart_search(content, search_term)

# ‡∏™‡∏≥‡∏´‡∏£‡∏±‡∏ö testing
if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        test_file = sys.argv[1]
    else:
        test_file = "test_document.pdf"
    
    if os.path.exists(test_file):
        print("Testing Enhanced Document Reader...")
        print("=" * 50)
        
        reader = EnhancedDocumentReader(test_file, use_ocr=True, expert_role="‡∏ú‡∏π‡πâ‡πÄ‡∏ä‡∏µ‡πà‡∏¢‡∏ß‡∏ä‡∏≤‡∏ç‡∏î‡πâ‡∏≤‡∏ô‡∏Å‡∏≤‡∏£‡∏®‡∏∂‡∏Å‡∏©‡∏≤")
        summary = reader.get_comprehensive_summary()
        
        print("Document Summary:")
        print(summary[:2000] + "..." if len(summary) > 2000 else summary)
        
        # Test search
        if len(sys.argv) > 2:
            search_term = sys.argv[2]
            print(f"\n\nSearching for '{search_term}':")
            print("=" * 50)
            search_result = reader.smart_search(summary, search_term)
            print(search_result)
    else:
        print(f"Test file '{test_file}' not found")
        print("Usage: python document_reader.py <file_path> [search_term]")